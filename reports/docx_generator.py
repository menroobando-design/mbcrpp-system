from docx import Document
from docx.shared import Cm
from django.conf import settings
import os

from .docx_helpers import add_photo_section


def replace_text(doc, old, new):
    new = str(new)

    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)

    # Replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)


def replace_logo(doc, placeholder, image_path):

    # Search normal paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:

                run.text = run.text.replace(
                    placeholder,
                    ""
                )

                run.add_picture(
                    image_path,
                    width=Cm(3),
                    height=Cm(3)
                )

                return

    # Search tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:

                            run.text = run.text.replace(
                                placeholder,
                                ""
                            )

                            run.add_picture(
                                image_path,
                                width=Cm(3),
                                height=Cm(3)
                            )

                            return


def generate_report_docx(report):

    template = os.path.join(
        settings.BASE_DIR,
        "report_templates",
        "Annex_A_Template.docx"
    )

    doc = Document(template)

    # ==========================
    # INSERT BARANGAY LOGO
    # ==========================

    if (
        report.barangay.logo
        and os.path.exists(report.barangay.logo.path)
    ):
        replace_logo(
            doc,
            "{{BARANGAY_LOGO}}",
            report.barangay.logo.path
        )

    # ==========================
    # REPLACE PLACEHOLDERS
    # ==========================

    replace_text(
        doc,
        "{{BARANGAY}}",
        report.barangay.name
    )

    replace_text(
        doc,
        "{{ACTIVITY_DATE}}",
        report.activity_date.strftime("%B %d, %Y")
    )

    replace_text(
        doc,
        "{{LOCATION}}",
        report.activity_location
    )

    replace_text(
        doc,
        "{{ TOTAL }}",
        f"{report.total_waste} kg"
    )

    replace_text(
        doc,
        "{{ DISPOSAL }}",
        report.disposal_method or ""
    )

    replace_text(
        doc,
        "{{ REMARKS }}",
        report.remarks or ""
    )

    replace_text(
        doc,
        "{{COMMITTEE_CHAIR}}",
        report.barangay.committee_chair or ""
    )

    replace_text(
        doc,
        "{{BARANGAY_CAPTAIN}}",
        report.barangay.barangay_captain or ""
    )

    # ==========================
    # PHOTO DOCUMENTATION
    # ==========================

    photo_sections = [

        ("Before", report.photos.filter(category="Before")),

        ("During", report.photos.filter(category="During")),

        ("After", report.photos.filter(category="After")),

        ("Collected Wastes", report.photos.filter(category="Collected Waste")),

        ("Group Photo", report.photos.filter(category="Group Photo")),

        ("Attendance", report.photos.filter(category="Attendance")),

    ]

    first_section = True

    for title, photos in photo_sections:

        if photos.exists():

            if not first_section:
                doc.add_page_break()

            add_photo_section(
                doc,
                title,
                photos
            )

            first_section = False

    # ==========================
    # SAVE DOCUMENT
    # ==========================

    filename = os.path.join(
        settings.BASE_DIR,
        "media",
        f"Report_{report.id}.docx"
    )

    doc.save(filename)

    return filename