import os

from PIL import Image

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


PHOTO_WIDTH = 15  # cm


def resize_image(image_path):
    img = Image.open(image_path)

    if img.mode != "RGB":
        img = img.convert("RGB")

    img.thumbnail((2500, 1800))

    base, ext = os.path.splitext(image_path)
    resized = base + "_resized.jpg"

    img.save(
        resized,
        "JPEG",
        quality=95
    )

    return resized


def add_category_title(document, title):

    heading = document.add_heading(
        title.upper(),
        level=1
    )

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_single_photo(document, photo):

    resized = resize_image(photo.image.path)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()

    # Width only
    # Height adjusts automatically
    run.add_picture(
        resized,
        width=Cm(PHOTO_WIDTH)
    )

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption.add_run(
        photo.caption or ""
    ).bold = True

    # Small spacing only
    document.add_paragraph()


def add_photo_section(document, title, photos):

    photos = list(photos)

    if not photos:
        return

    add_category_title(
        document,
        title
    )

    for photo in photos:
        add_single_photo(
            document,
            photo
        )